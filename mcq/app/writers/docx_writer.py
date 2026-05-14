"""DOCX writer.

Two output layouts:

* **normal**: 2-col question table (SL | question + A./B./C./D. options)
  followed by a 3-col answer sheet (Q No. | Ans | Explanation).

* **database**: single 8-col table — one row per question:
      (blank) | Question | OptA | OptB | OptC | OptD | (blank) | "Letter; Explanation"

Math handling:
  - Each text field is segmented via `math_utils.split_text`.
  - "text" segments → ordinary <w:r> runs.
  - "math" segments → pandoc-converted <m:oMath> elements injected as raw XML
    when pandoc is available, otherwise a Unicode-text fallback that keeps the
    KaTeX source visible (so nothing is silently dropped).
"""

from __future__ import annotations
import io
from typing import List

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor
from lxml import etree

from ..models import Question
from ..math_utils import split_text, katex_to_omml_xml


# ---------------------------------------------------------------------------
# Rich-text helper: writes a string into a paragraph, rendering inline KaTeX
# as real Word equations.
# ---------------------------------------------------------------------------

def _add_rich(paragraph, text: str, *, bold: bool = False, size_pt: int = 11):
    """Append `text` to `paragraph`, converting $...$ to native Word equations.

    Falls back to keeping the literal "$...$" source visible (no silent loss)
    if pandoc-based conversion fails for a particular expression.
    """
    if text is None:
        text = ""
    for kind, val in split_text(text):
        if kind == "text":
            if val:
                run = paragraph.add_run(val)
                run.bold = bold
                run.font.size = Pt(size_pt)
        else:  # math
            omml = katex_to_omml_xml(val)
            if omml:
                try:
                    el = etree.fromstring(omml)
                    paragraph._p.append(el)
                    continue
                except etree.XMLSyntaxError:
                    pass
            # Fallback: keep literal source so the data isn't lost.
            run = paragraph.add_run(f"${val}$")
            run.bold = bold
            run.font.size = Pt(size_pt)


def _set_cell_borders(cell, *, color: str = "999999"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = etree.SubElement(tcPr, qn("w:tcBorders"))
    for side in ("top", "left", "bottom", "right"):
        b = etree.SubElement(tcBorders, qn(f"w:{side}"))
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), color)


def _new_document(title: str | None = None) -> Document:
    doc = Document()
    # Page setup: A4, 1-inch margins. Matches both source docs.
    section = doc.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    if title:
        h = doc.add_paragraph()
        run = h.add_run(title)
        run.bold = True
        run.font.size = Pt(14)
    return doc


# ---------------------------------------------------------------------------
# Layout 1: Normal (question table + answer sheet)
# ---------------------------------------------------------------------------

_LETTERS = ["A", "B", "C", "D"]


def write_docx_normal(questions: List[Question],
                      *,
                      title: str = "Question Paper") -> bytes:
    doc = _new_document(title)

    # --- Question table ----------------------------------------------------
    qtbl = doc.add_table(rows=0, cols=2)
    qtbl.autofit = False
    # Approximate column widths (DXA via Cm helper).
    col_widths = [Cm(1.5), Cm(15)]
    for q in questions:
        row = qtbl.add_row()
        cells = row.cells
        # SL cell
        sl_p = cells[0].paragraphs[0]
        sl_run = sl_p.add_run(f"{q.sl:02d}.")
        sl_run.bold = True
        sl_run.font.size = Pt(11)
        cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP

        # Question + options cell
        qcell = cells[1]
        # First paragraph: question (bold)
        qp = qcell.paragraphs[0]
        _add_rich(qp, q.question, bold=True, size_pt=11)
        # Then options — each as its own paragraph.
        for idx, opt in enumerate(q.options):
            op = qcell.add_paragraph()
            op.paragraph_format.space_after = Pt(0)
            # "A." prefix
            prefix = op.add_run(f"{_LETTERS[idx]}. ")
            prefix.bold = True
            prefix.font.size = Pt(11)
            _add_rich(op, opt, bold=False, size_pt=11)

        _set_cell_borders(cells[0])
        _set_cell_borders(cells[1])
        # Set column widths
        cells[0].width = col_widths[0]
        cells[1].width = col_widths[1]

    # --- Answer sheet ------------------------------------------------------
    doc.add_paragraph()  # spacer
    head = doc.add_paragraph()
    run = head.add_run("Answer Sheet")
    run.bold = True
    run.font.size = Pt(13)

    atbl = doc.add_table(rows=1, cols=3)
    atbl.autofit = False
    hdr = atbl.rows[0].cells
    for i, label in enumerate(("Q No.", "Ans", "Explanation")):
        p = hdr[i].paragraphs[0]
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(11)
        _set_cell_borders(hdr[i])

    for q in questions:
        row = atbl.add_row()
        c = row.cells
        # Q No.
        c[0].paragraphs[0].add_run(f"{q.sl:02d}.").bold = False
        # Ans letter
        ap = c[1].paragraphs[0]
        ar = ap.add_run(q.answer_letter)
        ar.bold = True
        # Explanation (may contain math)
        ep = c[2].paragraphs[0]
        _add_rich(ep, q.explanation, bold=False, size_pt=10)
        for cell in c:
            _set_cell_borders(cell)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Layout 2: Database (8-column single table)
# ---------------------------------------------------------------------------

def write_docx_database(questions: List[Question],
                        *,
                        title: str | None = None) -> bytes:
    doc = _new_document(title)

    tbl = doc.add_table(rows=0, cols=8)
    tbl.autofit = True

    for q in questions:
        row = tbl.add_row()
        c = row.cells
        # 0: SL (a small index; useful for QA, kept short)
        c[0].paragraphs[0].add_run(f"{q.sl:02d}").bold = True
        # 1: question
        _add_rich(c[1].paragraphs[0], q.question, bold=True, size_pt=10)
        # 2..5: options
        for i in range(4):
            _add_rich(c[2 + i].paragraphs[0], q.options[i], size_pt=10)
        # 6: (intentionally blank — mirrors source format)
        # 7: "Letter; explanation"
        last = c[7].paragraphs[0]
        prefix = last.add_run(f"{q.answer_letter}; ")
        prefix.bold = True
        _add_rich(last, q.explanation, size_pt=10)
        for cell in c:
            _set_cell_borders(cell)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

"""DOCX writer — compact layout that mirrors the user's source paper.

Two output layouts:

**normal**: optional banner image + 2 tables.
  1. Question table: 2 cols × N rows.
       Col 0 = SL ("01."), narrow (~300 DXA).
       Col 1 = nested 4-column table identical to the source structure:
           Row 0  (gridSpan=4):  the question text, bold
           Row 1  (4 cells):     [A.] [option A] [B.] [option B]
           Row 2  (4 cells):     [C.] [option C] [D.] [option D]
       Marker cells are very narrow (~260 DXA) so "A." sits right next to
       the option text — no horizontal gap.
  2. Answer sheet — starts on a new page (via a NEW_PAGE section break).

**database**: optional banner image + single 8-column table:
       [SL] [Question] [optA] [optB] [optC] [optD] [Answer] [Explanation]
  (Source files leave column 7 blank; we emit the answer letter there.)

Every paragraph this writer emits has explicit zero spacing
(`spaceBefore=0`, `spaceAfter=0`, `line=240`/single) and zero indentation.

Optional `header_image` (bytes) is inserted at the very top of the body,
centred and sized to the page width. Pass `None` to skip.

Math handling (per `math_mode`):
  - "equation": $...$ KaTeX → real Word equations via OMML.
  - "text":     $...$ kept verbatim as text.
  - "unicode":  best-effort conversion (e.g. `x^2` → `x²`).
"""

from __future__ import annotations
import io
from typing import List, Optional

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, Emu, Inches
from docx.text.paragraph import Paragraph
from lxml import etree

from ..models import Question
from ..math_utils import split_text, katex_to_omml_xml, katex_to_unicode


_W_NS_DECL = (
    'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
)


# ---------------------------------------------------------------------------
# Paragraph-level helpers
# ---------------------------------------------------------------------------

def _zero_paragraph_spacing(p_el):
    """Set spaceBefore=0, spaceAfter=0, line=240 (single), ind=0 on a <w:p>.

    Word's default style adds ~8pt of space after every paragraph and runs
    line height at ~1.08. Both are why our output had big vertical gaps
    between the question and the options. We zero them out explicitly on
    every paragraph we emit so the layout matches the source's tight stack.
    """
    # Find or create the <w:pPr> child.
    pPr = p_el.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        # pPr must be the first child of <w:p>
        p_el.insert(0, pPr)

    # Drop any existing spacing/ind so we own them.
    for old in pPr.findall(qn("w:spacing")):
        pPr.remove(old)
    for old in pPr.findall(qn("w:ind")):
        pPr.remove(old)
    for old in pPr.findall(qn("w:contextualSpacing")):
        pPr.remove(old)

    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), "0")
    sp.set(qn("w:after"), "0")
    sp.set(qn("w:line"), "240")          # 240 twentieths = single line
    sp.set(qn("w:lineRule"), "auto")
    pPr.append(sp)

    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "0")
    ind.set(qn("w:right"), "0")
    ind.set(qn("w:firstLine"), "0")
    pPr.append(ind)

    cs = OxmlElement("w:contextualSpacing")
    pPr.append(cs)


def _add_rich(paragraph, text: str, *,
              bold: bool = False, size_pt: int = 10,
              math_mode: str = "equation"):
    """Append `text` to `paragraph`, handling $...$ math per math_mode.

    Also zeroes the paragraph's spacing so callers don't have to remember.
    """
    if text is None:
        text = ""
    if math_mode not in ("equation", "text", "unicode"):
        raise ValueError(f"unknown math_mode: {math_mode!r}")

    _zero_paragraph_spacing(paragraph._p)

    for kind, val in split_text(text):
        if kind == "text":
            if val:
                run = paragraph.add_run(val)
                run.bold = bold
                run.font.size = Pt(size_pt)
            continue
        # math segment
        if math_mode == "equation":
            omml = katex_to_omml_xml(val)
            if omml:
                try:
                    el = etree.fromstring(omml)
                    paragraph._p.append(el)
                    continue
                except etree.XMLSyntaxError:
                    pass
            literal = f"${val}$"
        elif math_mode == "unicode":
            literal = katex_to_unicode(val)
        else:  # text
            literal = f"${val}$"
        run = paragraph.add_run(literal)
        run.bold = bold
        run.font.size = Pt(size_pt)


def _fill_paragraph_xml(p_el, text: str, *, bold: bool = False,
                        size_pt: int = 10, math_mode: str = "equation"):
    """Fill an existing raw <w:p> element with text+math, zeroing spacing."""
    para = Paragraph(p_el, parent=None)
    _add_rich(para, text, bold=bold, size_pt=size_pt, math_mode=math_mode)


def _fill_paragraph_with_sl_and_question(p_el, *, sl: int, question: str,
                                         math_mode: str = "equation"):
    """Write "NN. {question}" into a paragraph with the SL bold and inline.

    The SL is bold and slightly larger; the question text is bold but
    normal size. They sit on the same line — when the question wraps to
    a second line, Word handles the hang naturally because there's no
    indentation set.
    """
    para = Paragraph(p_el, parent=None)
    _zero_paragraph_spacing(p_el)

    # SL prefix — bold, same size as question
    sl_run = para.add_run(f"{sl:02d}. ")
    sl_run.bold = True
    sl_run.font.size = Pt(10)

    # Question text with math, also bold
    from ..math_utils import split_text, katex_to_omml_xml, katex_to_unicode
    for kind, val in split_text(question or ""):
        if kind == "text":
            if val:
                r = para.add_run(val)
                r.bold = True
                r.font.size = Pt(10)
            continue
        # math
        if math_mode == "equation":
            omml = katex_to_omml_xml(val)
            if omml:
                try:
                    el = etree.fromstring(omml)
                    para._p.append(el)
                    continue
                except etree.XMLSyntaxError:
                    pass
            literal = f"${val}$"
        elif math_mode == "unicode":
            literal = katex_to_unicode(val)
        else:
            literal = f"${val}$"
        r = para.add_run(literal)
        r.bold = True
        r.font.size = Pt(10)


def _set_cell_borders(cell, *, color: str = "cccccc"):
    tcPr = cell._tc.get_or_add_tcPr()
    # Drop existing
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    tcBorders = etree.SubElement(tcPr, qn("w:tcBorders"))
    for side in ("top", "left", "bottom", "right"):
        b = etree.SubElement(tcBorders, qn(f"w:{side}"))
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), color)


def _new_document(title: str | None = None) -> Document:
    """Create a blank document with sensible page margins.

    The `title` argument is accepted for backward compatibility but no
    longer rendered here — use `_add_paper_title` after the header image
    so the title appears centered just below the banner.
    """
    doc = Document()
    s = doc.sections[0]
    s.left_margin = Cm(1.5)
    s.right_margin = Cm(1.5)
    s.top_margin = Cm(1.5)
    s.bottom_margin = Cm(1.5)
    return doc


def _add_paper_title(doc: Document, title: str) -> None:
    """Insert the paper title (display name + 'Set N') centered & bold.

    Goes immediately after the header banner (or at the very top if there
    is no banner). The title sits at the centre-bottom of the header area
    so students see the paper name and their set number together at the
    top of page 1, before the questions begin.
    """
    if not title:
        return
    h = doc.add_paragraph()
    _zero_paragraph_spacing(h._p)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # A small space before so the title doesn't collide with the banner
    h.paragraph_format.space_before = Pt(2)
    h.paragraph_format.space_after = Pt(4)
    r = h.add_run(title)
    r.bold = True
    r.font.size = Pt(13)


def _set_section_columns(section, num_cols: int, space_dxa: int = 360):
    sectPr = section._sectPr
    for old in sectPr.findall(qn("w:cols")):
        sectPr.remove(old)
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), str(num_cols))
    cols.set(qn("w:space"), str(space_dxa))
    sectPr.append(cols)


def _start_new_section(doc: Document, *, num_cols: int,
                       new_page: bool = False) -> None:
    section_type = WD_SECTION.NEW_PAGE if new_page else WD_SECTION.CONTINUOUS
    new = doc.add_section(section_type)
    new.left_margin = Cm(1.5)
    new.right_margin = Cm(1.5)
    _set_section_columns(new, num_cols)


def _add_header_image(doc: Document, image_bytes: bytes) -> None:
    """Insert a centred image at the very top of the body.

    Sized to the page's usable width (about 18 cm at 1.5 cm margins on A4).
    Used for the question paper banner — appears once at the top of page 1.
    """
    p = doc.add_paragraph()
    _zero_paragraph_spacing(p._p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    # Fit width: A4 = 21 cm, minus 1.5 cm × 2 margins = 18 cm usable.
    run.add_picture(io.BytesIO(image_bytes), width=Cm(18))


def _set_section_columns_only(section, num_cols: int, space_dxa: int = 360):
    """Public alias kept for clarity in the layout functions."""
    _set_section_columns(section, num_cols, space_dxa)


# ---------------------------------------------------------------------------
# Layout 1: Normal
# ---------------------------------------------------------------------------

def _nested_options_table_xml() -> str:
    """The inner 4-column table that holds the question (row 0, gridSpan=4)
    and the options (rows 1 & 2, with narrow marker cells).

    Each row has <w:cantSplit/> so it never splits across a column/page
    boundary, and the question + first option row carry <w:keepNext/> in
    their paragraph properties — together these guarantee the whole
    question (label + 4 options) stays in a single column.
    """
    return f'''
<w:tbl {_W_NS_DECL}>
  <w:tblPr>
    <w:tblW w:w="5000" w:type="pct"/>
    <w:tblBorders>
      <w:top w:val="nil"/><w:left w:val="nil"/>
      <w:bottom w:val="nil"/><w:right w:val="nil"/>
      <w:insideH w:val="nil"/><w:insideV w:val="nil"/>
    </w:tblBorders>
    <w:tblLayout w:type="fixed"/>
    <w:tblCellMar>
      <w:top w:w="0" w:type="dxa"/>
      <w:left w:w="40" w:type="dxa"/>
      <w:bottom w:w="0" w:type="dxa"/>
      <w:right w:w="40" w:type="dxa"/>
    </w:tblCellMar>
  </w:tblPr>
  <w:tblGrid>
    <w:gridCol w:w="260"/>
    <w:gridCol w:w="1940"/>
    <w:gridCol w:w="260"/>
    <w:gridCol w:w="1940"/>
  </w:tblGrid>
  <w:tr>
    <w:trPr><w:cantSplit/></w:trPr>
    <w:tc>
      <w:tcPr><w:tcW w:w="4400" w:type="dxa"/><w:gridSpan w:val="4"/></w:tcPr>
      <w:p><w:pPr><w:keepNext/><w:keepLines/></w:pPr></w:p>
    </w:tc>
  </w:tr>
  <w:tr>
    <w:trPr><w:cantSplit/></w:trPr>
    <w:tc><w:tcPr><w:tcW w:w="260"  w:type="dxa"/></w:tcPr><w:p><w:pPr><w:keepNext/></w:pPr></w:p></w:tc>
    <w:tc><w:tcPr><w:tcW w:w="1940" w:type="dxa"/></w:tcPr><w:p><w:pPr><w:keepNext/></w:pPr></w:p></w:tc>
    <w:tc><w:tcPr><w:tcW w:w="260"  w:type="dxa"/></w:tcPr><w:p><w:pPr><w:keepNext/></w:pPr></w:p></w:tc>
    <w:tc><w:tcPr><w:tcW w:w="1940" w:type="dxa"/></w:tcPr><w:p><w:pPr><w:keepNext/></w:pPr></w:p></w:tc>
  </w:tr>
  <w:tr>
    <w:trPr><w:cantSplit/></w:trPr>
    <w:tc><w:tcPr><w:tcW w:w="260"  w:type="dxa"/></w:tcPr><w:p><w:pPr><w:keepLines/></w:pPr></w:p></w:tc>
    <w:tc><w:tcPr><w:tcW w:w="1940" w:type="dxa"/></w:tcPr><w:p><w:pPr><w:keepLines/></w:pPr></w:p></w:tc>
    <w:tc><w:tcPr><w:tcW w:w="260"  w:type="dxa"/></w:tcPr><w:p><w:pPr><w:keepLines/></w:pPr></w:p></w:tc>
    <w:tc><w:tcPr><w:tcW w:w="1940" w:type="dxa"/></w:tcPr><w:p><w:pPr><w:keepLines/></w:pPr></w:p></w:tc>
  </w:tr>
</w:tbl>'''.strip()


def write_docx_normal(questions: List[Question],
                      *,
                      title: str = "",
                      math_mode: str = "equation",
                      header_image: Optional[bytes] = None) -> bytes:
    doc = _new_document(title)

    # Header banner (page 1 only — it's body content, not a Word "header")
    if header_image:
        # Insert into the FIRST section, which is single-column by default.
        # The 2-column layout for questions is applied via a continuous section
        # break right after, so the header spans the full page width.
        try:
            _add_header_image(doc, header_image)
            # Title goes inside/just below the banner, centred + bold.
            _add_paper_title(doc, title)
            _start_new_section(doc, num_cols=2, new_page=False)
        except Exception:
            # Bad image bytes shouldn't kill the whole generation.
            _add_paper_title(doc, title)
            _set_section_columns(doc.sections[0], num_cols=2, space_dxa=300)
    else:
        # No header → put title at the very top, then switch to 2-column
        _add_paper_title(doc, title)
        _start_new_section(doc, num_cols=2, new_page=False)

    body = doc.element.body

    for q in questions:
        # One nested-style table per question — three rows, no outer wrapper.
        # Row 0 (gridSpan=4): "NN. {question text}" with SL inline so it's
        # guaranteed on the same line as the start of the question.
        # Rows 1 & 2: option grid with narrow 260-DXA marker cells.
        tbl = etree.fromstring(_nested_options_table_xml())
        sectPr = body.find(qn("w:sectPr"))
        if sectPr is not None:
            sectPr.addprevious(tbl)
        else:
            body.append(tbl)

        ntrs = tbl.findall(qn("w:tr"))

        # Row 0 — question (with inline SL)
        q_tc = ntrs[0].find(qn("w:tc"))
        q_p = q_tc.find(qn("w:p"))
        _fill_paragraph_with_sl_and_question(
            q_p, sl=q.sl, question=q.question,
            math_mode=math_mode,
        )

        # Row 1 — A / B (markers bold, option text NOT bold)
        tcs = ntrs[1].findall(qn("w:tc"))
        _fill_paragraph_xml(tcs[0].find(qn("w:p")), "A.", bold=True,
                            size_pt=10, math_mode=math_mode)
        _fill_paragraph_xml(tcs[1].find(qn("w:p")), q.options[0],
                            bold=False, size_pt=10, math_mode=math_mode)
        _fill_paragraph_xml(tcs[2].find(qn("w:p")), "B.", bold=True,
                            size_pt=10, math_mode=math_mode)
        _fill_paragraph_xml(tcs[3].find(qn("w:p")), q.options[1],
                            bold=False, size_pt=10, math_mode=math_mode)

        # Row 2 — C / D
        tcs = ntrs[2].findall(qn("w:tc"))
        _fill_paragraph_xml(tcs[0].find(qn("w:p")), "C.", bold=True,
                            size_pt=10, math_mode=math_mode)
        _fill_paragraph_xml(tcs[1].find(qn("w:p")), q.options[2],
                            bold=False, size_pt=10, math_mode=math_mode)
        _fill_paragraph_xml(tcs[2].find(qn("w:p")), "D.", bold=True,
                            size_pt=10, math_mode=math_mode)
        _fill_paragraph_xml(tcs[3].find(qn("w:p")), q.options[3],
                            bold=False, size_pt=10, math_mode=math_mode)

    # Answer sheet — starts on a NEW PAGE, single-column, full width.
    _start_new_section(doc, num_cols=1, new_page=True)

    head = doc.add_paragraph()
    _zero_paragraph_spacing(head._p)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = head.add_run("Answer Sheet")
    hr.bold = True
    hr.font.size = Pt(14)

    # Blank line for breathing room
    spacer = doc.add_paragraph()
    _zero_paragraph_spacing(spacer._p)

    atbl = doc.add_table(rows=1, cols=3)
    atbl.autofit = False
    hdr = atbl.rows[0].cells
    for i, label in enumerate(("Q No.", "Ans", "Explanation")):
        p = hdr[i].paragraphs[0]
        _zero_paragraph_spacing(p._p)
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(10)
        _set_cell_borders(hdr[i])

    for q in questions:
        row = atbl.add_row()
        c = row.cells
        p0 = c[0].paragraphs[0]
        _zero_paragraph_spacing(p0._p)
        p0.add_run(f"{q.sl:02d}.").font.size = Pt(10)

        p1 = c[1].paragraphs[0]
        _zero_paragraph_spacing(p1._p)
        ar = p1.add_run(q.answer_letter)
        ar.bold = True
        ar.font.size = Pt(10)

        _add_rich(c[2].paragraphs[0], q.explanation,
                  bold=False, size_pt=10, math_mode=math_mode)
        for cell in c:
            _set_cell_borders(cell)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Layout 2: Database (8-column single table) — compact spacing as well
# ---------------------------------------------------------------------------

def write_docx_database(questions: List[Question],
                        *,
                        title: str | None = None,
                        math_mode: str = "equation",
                        header_image: Optional[bytes] = None) -> bytes:
    """Single 8-column table — one row per question.

    Column layout:
        0: SL      1: Question     2-5: Options A-D
        6: Answer letter (A/B/C/D)  ← was blank in your source; now populated
        7: Explanation (without redundant leading letter)
    """
    doc = _new_document(title)

    if header_image:
        try:
            _add_header_image(doc, header_image)
        except Exception:
            pass

    # Centered, bold paper title — goes below the banner if there is one,
    # else at the top of the page.
    _add_paper_title(doc, title or "")

    tbl = doc.add_table(rows=0, cols=8)
    tbl.autofit = True

    for q in questions:
        row = tbl.add_row()
        c = row.cells

        # SL
        p = c[0].paragraphs[0]
        _zero_paragraph_spacing(p._p)
        r = p.add_run(f"{q.sl:02d}")
        r.bold = True
        r.font.size = Pt(10)

        # Question — bold
        _add_rich(c[1].paragraphs[0], q.question,
                  bold=True, size_pt=10, math_mode=math_mode)

        # Options A-D — not bold
        for i in range(4):
            _add_rich(c[2 + i].paragraphs[0], q.options[i],
                      bold=False, size_pt=10, math_mode=math_mode)

        # Column 7 (index 6): the answer letter — was blank in source files,
        # now contains the correct answer so the table is self-explanatory.
        ap = c[6].paragraphs[0]
        _zero_paragraph_spacing(ap._p)
        ar = ap.add_run(q.answer_letter)
        ar.bold = True
        ar.font.size = Pt(10)

        # Column 8 (index 7): explanation only — no letter prefix any more
        # (the dedicated answer column above has it).
        _add_rich(c[7].paragraphs[0], q.explanation,
                  bold=False, size_pt=10, math_mode=math_mode)

        for cell in c:
            _set_cell_borders(cell)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
